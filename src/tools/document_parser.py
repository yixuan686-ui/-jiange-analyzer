"""
文档解析工具
支持多种文件格式的解析，包括PDF、Word、Excel、文本文件等
"""

from langchain.tools import tool, ToolRuntime
from coze_coding_utils.runtime_ctx.context import new_context
from typing import Literal
import os
import json
import csv

# 尝试导入文件解析库，如果不存在则提示安装
PDF_AVAILABLE = False
DOCX_AVAILABLE = False
EXCEL_AVAILABLE = False

try:
    from pypdf import PdfReader as pypdf_PdfReader  # type: ignore
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf  # type: ignore
        PDF_AVAILABLE = True
    except ImportError:
        pass

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import openpyxl  # type: ignore
    EXCEL_AVAILABLE = True
except ImportError:
    pass


def read_text_file(file_path: str) -> str:
    """读取文本文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试其他编码
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except:
            return "无法读取文件：编码不支持"


def read_pdf_file(file_path: str) -> str:
    """读取PDF文件"""
    if not PDF_AVAILABLE:
        return "PDF解析功能未安装，请运行: pip install pypdf"
    
    try:
        text_parts = []
        with open(file_path, 'rb') as f:
            try:
                reader = pypdf_PdfReader(f)
            except:
                import pypdf
                reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text())
        return '\n\n'.join(text_parts)
    except Exception as e:
        return f"PDF解析失败: {str(e)}"


def read_docx_file(file_path: str) -> str:
    """读取Word文件"""
    if not DOCX_AVAILABLE:
        return "Word解析功能未安装，请运行: pip install python-docx"
    
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        return '\n'.join(text_parts)
    except Exception as e:
        return f"Word解析失败: {str(e)}"


def read_excel_file(file_path: str) -> str:
    """读取Excel文件"""
    if not EXCEL_AVAILABLE:
        return "Excel解析功能未安装，请运行: pip install openpyxl"
    
    try:
        import openpyxl
        workbook = openpyxl.load_workbook(file_path)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"\n=== 工作表: {sheet_name} ===\n")
            
            for row in sheet.iter_rows(values_only=True):
                # 过滤空行
                if any(cell is not None for cell in row):
                    # 将单元格值转换为字符串并用制表符分隔
                    row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                    text_parts.append(row_text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        return f"Excel解析失败: {str(e)}"


def read_csv_file(file_path: str) -> str:
    """读取CSV文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = list(reader)
            return '\n'.join(','.join(row) for row in rows)
    except Exception as e:
        return f"CSV解析失败: {str(e)}"


def read_json_file(file_path: str) -> str:
    """读取JSON文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json.dumps(data, ensure_ascii=False, indent=2)
    except Exception as e:
        return f"JSON解析失败: {str(e)}"


@tool
def parse_document(
    file_path: str,
    file_type: Literal["auto", "txt", "md", "csv", "json", "pdf", "docx", "xlsx"] = "auto",
    runtime: ToolRuntime = None
) -> str:
    """
    解析文档文件内容，支持多种格式
    
    Args:
        file_path: 文件路径（可以是绝对路径或相对于assets目录的路径）
        file_type: 文件类型，"auto"表示自动根据文件扩展名推断
                   可选值: auto, txt, md, csv, json, pdf, docx, xlsx
        runtime: 工具运行时上下文
    
    Returns:
        解析后的文本内容
    """
    # 如果是相对路径，尝试在assets目录中查找
    if not os.path.isabs(file_path):
        workspace_path = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
        assets_path = os.path.join(workspace_path, "assets", file_path)
        if os.path.exists(assets_path):
            file_path = assets_path
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"
    
    # 自动推断文件类型
    if file_type == "auto":
        ext = os.path.splitext(file_path)[1].lower()
        ext_map = {
            '.txt': 'txt',
            '.md': 'md',
            '.csv': 'csv',
            '.json': 'json',
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.xlsx': 'xlsx'
        }
        file_type = ext_map.get(ext, 'txt')
    
    # 根据文件类型选择解析方法
    parser_map = {
        'txt': read_text_file,
        'md': read_text_file,
        'csv': read_csv_file,
        'json': read_json_file,
        'pdf': read_pdf_file,
        'docx': read_docx_file,
        'xlsx': read_excel_file
    }
    
    parser = parser_map.get(file_type, read_text_file)
    
    try:
        content = parser(file_path)
        return f"文件类型: {file_type}\n文件路径: {file_path}\n\n--- 文件内容 ---\n\n{content}"
    except Exception as e:
        return f"文件解析失败: {str(e)}"


@tool
def parse_url_content(
    url: str,
    runtime: ToolRuntime = None
) -> str:
    """
    解析网页URL内容（使用知识库SDK的URL导入功能）
    
    Args:
        url: 网页URL
        runtime: 工具运行时上下文
    
    Returns:
        网页内容的摘要和关键信息
    """
    try:
        from coze_coding_dev_sdk import KnowledgeClient, Config, KnowledgeDocument, DataSourceType
        
        ctx = runtime.context if runtime else new_context(method="url_parse")
        config = Config()
        client = KnowledgeClient(config=config, ctx=ctx)
        
        # 导入URL到知识库
        doc = KnowledgeDocument(
            source=DataSourceType.URL,
            url=url
        )
        
        response = client.add_documents(
            documents=[doc],
            table_name="temp_url_parse"
        )
        
        if response.code != 0:
            return f"网址解析失败: {response.msg}"
        
        # 搜索刚导入的内容
        search_response = client.search(
            query="网页主要内容",
            table_names=["temp_url_parse"],
            top_k=3
        )
        
        if search_response.code != 0:
            return f"内容提取失败: {search_response.msg}"
        
        # 整理结果
        result_parts = [f"URL: {url}", "\n--- 网页内容摘要 ---\n"]
        
        for chunk in search_response.chunks:
            result_parts.append(chunk.content)
        
        return '\n\n'.join(result_parts)
        
    except Exception as e:
        return f"网址解析失败: {str(e)}"
